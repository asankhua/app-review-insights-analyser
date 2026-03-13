"""
Phase 4 Email Delivery Service
"""
import json
import logging
import re
import time
from datetime import datetime, date
from typing import List, Dict, Any, Optional
from pathlib import Path

from .models.email import (
    EmailMessage, 
    EmailDeliveryRequest, 
    EmailDeliveryResponse, 
    EmailDeliveryFile,
    EmailMode,
    EmailStatus,
    EmailAttachment
)
from .email_service import EmailService
from phase3.note_generation import WeeklyNoteService
from .config.email_templates import (
    WEEKLY_PULSE_TEMPLATE,
    HTML_TEMPLATE,
    PLAIN_TEXT_TEMPLATE,
    format_markdown_to_html,
    strip_markdown_headers,
)
from src.config.settings import Config

logger = logging.getLogger(__name__)

class EmailDeliveryService:
    """Service for Phase 4: Email Delivery"""
    
    def __init__(self):
        self.config = Config()
        self.email_service = EmailService()
        self.note_service = WeeklyNoteService()
        
        # Ensure directories exist
        Path("data/drafts").mkdir(exist_ok=True)
        Path("data/deliveries").mkdir(exist_ok=True)
    
    def deliver_weekly_note(
        self,
        weekly_note_path: Optional[str] = None,
        weekly_note_content: Optional[str] = None,
        recipient_email: Optional[str] = None,
        recipient_name: Optional[str] = None,
        mode: EmailMode = EmailMode.DRY_RUN,
        include_attachments: bool = True,
        custom_subject: Optional[str] = None
    ) -> EmailDeliveryResponse:
        """
        Deliver weekly note via email
        
        Args:
            weekly_note_path: Path to weekly note file (auto-detect if None)
            weekly_note_content: Raw report content (e.g. from Gist). When set, overrides path.
            recipient_email: Override recipient email
            recipient_name: Recipient name for greeting
            mode: Email delivery mode (dry_run or send)
            include_attachments: Whether to include report attachments
            custom_subject: Custom email subject
            
        Returns:
            EmailDeliveryResponse with delivery results
        """
        try:
            logger.info(f"Starting email delivery: mode={mode}")
            start_time = time.time()
            
            # Prefer content (from Gist) over path (files)
            if not weekly_note_content:
                if not weekly_note_path:
                    weekly_note_path = self._get_latest_weekly_note_path()
                if not weekly_note_path or not Path(weekly_note_path).exists():
                    raise ValueError("No weekly note file found. Run Phase 3 first.")
                weekly_note_content = self._load_weekly_note_content(weekly_note_path)
            else:
                weekly_note_path = ""  # No file path when content from Gist
            
            # Determine recipient
            recipient = recipient_email or self.config.EMAIL_RECIPIENT or "team@indmoney.com"
            recipient_display_name = recipient_name or "Team"
            
            # Generate email content (pass content for attachment when path empty, e.g. from Gist)
            email_message = self._create_email_message(
                weekly_note_content=weekly_note_content,
                recipient_email=recipient,
                recipient_name=recipient_display_name,
                custom_subject=custom_subject,
                include_attachments=include_attachments,
                weekly_note_path=weekly_note_path,
                weekly_note_content_for_attach=weekly_note_content if not weekly_note_path else None,
            )
            
            # Send email
            delivery_result = self.email_service.send_email(email_message, mode.value)
            
            # Create response
            response = EmailDeliveryResponse(
                message_id=delivery_result['message_id'],
                status=delivery_result['status'],
                sent_at=datetime.fromisoformat(delivery_result['sent_at']) if delivery_result['sent_at'] else None,
                error_message=delivery_result.get('error_message'),
                draft_path=delivery_result.get('draft_path'),
                recipient=recipient,
                subject=email_message.subject,
                processing_time=time.time() - start_time
            )
            
            # Save delivery record (use placeholder path when content from Gist, for model validation)
            record_path = weekly_note_path or "INDMoney_Weekly_Pulse.md"
            self._save_delivery_record(
                request=EmailDeliveryRequest(
                    weekly_note_path=record_path,
                    recipient_email=recipient_email,
                    recipient_name=recipient_name,
                    mode=mode,
                    include_attachments=include_attachments,
                    custom_subject=custom_subject
                ),
                response=response
            )
            
            logger.info(f"Email delivery completed: {response.status.value}")
            return response
            
        except Exception as e:
            logger.error(f"Email delivery failed: {str(e)}")
            raise
    
    def _get_latest_weekly_note_path(self) -> Optional[str]:
        """Get path to latest weekly note file. Falls back to sample_data when no report exists."""
        try:
            latest_note = self.note_service.get_latest_weekly_note_file()
            if latest_note:
                path = f"data/reports/pulse-{latest_note.weekStart.strftime('%Y-%m-%d')}.md"
                if Path(path).exists():
                    return path
            return self._get_report_fallback_path()
        except Exception as e:
            logger.error(f"Error getting latest weekly note: {str(e)}")
            return self._get_report_fallback_path()

    def _get_report_fallback_path(self) -> Optional[str]:
        """Fallback: check data/reports and sample_data for pulse-*.md (same as pipeline)."""
        project_root = Path(__file__).parent.parent
        for dir_name in ("data/reports", "sample_data"):
            dir_path = project_root / dir_name
            if dir_path.exists():
                pulse_files = list(dir_path.glob("pulse-*.md"))
                if pulse_files:
                    return str(sorted(pulse_files)[-1])
        return None
    
    def _load_weekly_note_content(self, weekly_note_path: str) -> str:
        """Load weekly note content from file"""
        try:
            path = Path(weekly_note_path)
            
            if weekly_note_path.endswith('.md'):
                return path.read_text(encoding='utf-8')
            elif weekly_note_path.endswith('.json'):
                # Extract content from JSON file
                data = json.loads(path.read_text(encoding='utf-8'))
                return data.get('content', '')
            else:
                raise ValueError(f"Unsupported file format: {weekly_note_path}")
                
        except Exception as e:
            logger.error(f"Error loading weekly note: {str(e)}")
            raise
    
    def _create_email_message(
        self,
        weekly_note_content: str,
        recipient_email: str,
        recipient_name: str,
        custom_subject: Optional[str],
        include_attachments: bool,
        weekly_note_path: str,
        weekly_note_content_for_attach: Optional[str] = None,
    ) -> EmailMessage:
        """Create email message with formatted content"""
        try:
            # Extract week date from content
            week_date = self._extract_week_date(weekly_note_content)
            
            # Create email subject
            subject = custom_subject or f"INDMoney Weekly Review Pulse -- {week_date}"
            
            # Snippet for appended doc (first ~200 chars, strip markdown for display)
            raw_snippet = strip_markdown_headers(weekly_note_content).strip()[:250].replace("\n", " ")
            note_snippet = re.sub(r"\s+", " ", raw_snippet).strip()
            if len(strip_markdown_headers(weekly_note_content).strip()) > 250:
                note_snippet += "..."
            
            # Attach filename (always .docx)
            attach_filename = (
                Path(weekly_note_path).stem + ".docx"
                if weekly_note_path and Path(weekly_note_path).exists() and weekly_note_path.endswith(".md")
                else self._make_attachment_filename(week_date)
            )
            
            # Create DOCX bytes for attachment (doc attached to mail, no link)
            docx_bytes = None
            if weekly_note_content_for_attach or (weekly_note_path and Path(weekly_note_path).exists() and weekly_note_path.endswith(".md")):
                content_for_docx = weekly_note_content_for_attach or Path(weekly_note_path).read_text(encoding="utf-8")
                docx_bytes = self._markdown_to_docx(content_for_docx)
            
            # Generate email body (strip # headers to avoid duplicates)
            note_text = strip_markdown_headers(weekly_note_content)
            text_body = PLAIN_TEXT_TEMPLATE.format(
                recipient_name=recipient_name,
                week_date=week_date,
                weekly_note_text=note_text,
                generated_date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                appended_snippet=note_snippet,
                appended_filename=attach_filename,
            )
            
            # Convert markdown to HTML (headers stripped inside)
            weekly_note_html = format_markdown_to_html(weekly_note_content)
            html_body = HTML_TEMPLATE.format(
                recipient_name=recipient_name,
                week_date=week_date,
                weekly_note_html=weekly_note_html,
                generated_date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                appended_snippet=note_snippet,
                appended_filename=attach_filename,
            )
            
            # Create attachments (from file path or docx bytes)
            attachments = []
            if include_attachments:
                attachments = self._create_attachments(
                    weekly_note_path,
                    weekly_note_content_for_attach,
                    attach_filename,
                    docx_bytes,
                )
            
            return EmailMessage(
                subject=subject,
                from_email=self.config.EMAIL_SENDER,
                to_email=recipient_email,
                html_body=html_body,
                text_body=text_body,
                attachments=attachments
            )
            
        except Exception as e:
            logger.error(f"Error creating email message: {str(e)}")
            raise
    
    def _make_attachment_filename(self, week_date: str) -> str:
        """Create attachment filename from week date, e.g. INDMoney_Weekly_Pulse_March09-Mar15.docx"""
        safe = re.sub(r"[^\w\s-]", "", week_date).strip().replace(" ", "_")[:40]
        return f"INDMoney_Weekly_Pulse_{safe}.docx"

    def _markdown_to_docx(self, content: str) -> bytes:
        """Convert markdown content to Word .docx bytes."""
        try:
            from docx import Document
            doc = Document()
            for line in content.split("\n"):
                line = line.rstrip()
                if not line:
                    continue
                if line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=0)
                elif line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=1)
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
                elif line.strip().startswith("Action ") and ":" in line:
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph(line.strip())
            import io
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as e:
            logger.warning("Markdown to DOCX failed, falling back to plain text: %s", e)
            from docx import Document
            doc = Document()
            doc.add_paragraph(content)
            import io
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

    def _extract_week_date(self, content: str) -> str:
        """Extract week date range from weekly note content (e.g. 'March 2 - March 8')"""
        try:
            lines = content.split('\n')
            for line in lines:
                # Format: "## ... -- March 2 - March 8" or "Week of March 2 - March 8"
                if ' -- ' in line:
                    return line.split(' -- ')[-1].strip()
                if 'Week of' in line:
                    return line.split('Week of')[-1].strip()
            return "Unknown"
        except Exception:
            return "Unknown"
    
    def _create_attachments(
        self,
        weekly_note_path: str,
        weekly_note_content: Optional[str] = None,
        attach_filename: Optional[str] = None,
        docx_bytes: Optional[bytes] = None,
    ) -> List[EmailAttachment]:
        """Create email attachments from file path, content, or pre-built docx bytes."""
        attachments = []
        try:
            # Use pre-built docx if provided (from Drive upload path)
            if docx_bytes and attach_filename:
                attachment = EmailAttachment(
                    filename=attach_filename,
                    content=docx_bytes,
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    size=len(docx_bytes),
                )
                attachments.append(attachment)
            # From file path: convert .md to .docx (when docx_bytes not pre-built)
            elif weekly_note_path and weekly_note_path.strip() and not docx_bytes:
                path = Path(weekly_note_path)
                if path.exists() and path.suffix.lower() == ".md":
                    md_text = path.read_text(encoding="utf-8")
                    docx_bytes = self._markdown_to_docx(md_text)
                    filename = path.stem + ".docx"
                    attachment = EmailAttachment(
                        filename=filename,
                        content=docx_bytes,
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        size=len(docx_bytes),
                    )
                    attachments.append(attachment)
            # From content (e.g. Gist): convert to .docx (when docx_bytes not pre-built)
            elif weekly_note_content and attach_filename and not docx_bytes:
                docx_bytes = self._markdown_to_docx(weekly_note_content)
                attachment = EmailAttachment(
                    filename=attach_filename,
                    content=docx_bytes,
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    size=len(docx_bytes),
                )
                attachments.append(attachment)
        except Exception as e:
            logger.error(f"Error creating attachments: {str(e)}")
        return attachments
    
    def _save_delivery_record(
        self, 
        request: EmailDeliveryRequest, 
        response: EmailDeliveryResponse
    ) -> None:
        """Save delivery record to file"""
        try:
            delivery_file = EmailDeliveryFile(
                message_id=response.message_id,
                created_at=datetime.now(),
                request=request,
                response=response,
                metadata={
                    'processing_time': response.processing_time,
                    'email_service_config': self.email_service.get_configuration_info()
                }
            )
            
            # Save to deliveries directory
            deliveries_dir = Path("data/deliveries")
            filename = f"delivery_{response.message_id}.json"
            filepath = deliveries_dir / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(delivery_file.dict(), f, ensure_ascii=False, indent=2, default=str)
            
            logger.info(f"Delivery record saved: {filepath}")
            
        except Exception as e:
            logger.error(f"Error saving delivery record: {str(e)}")
    
    def get_delivery_stats(self) -> Dict[str, Any]:
        """Get email delivery statistics"""
        try:
            deliveries_dir = Path("data/deliveries")
            
            if not deliveries_dir.exists():
                return {
                    'total_sent': 0,
                    'total_drafts': 0,
                    'last_sent': None,
                    'success_rate': 0.0,
                    'most_recent_file': None
                }
            
            delivery_files = list(deliveries_dir.glob("delivery_*.json"))
            
            if not delivery_files:
                return {
                    'total_sent': 0,
                    'total_drafts': 0,
                    'last_sent': None,
                    'success_rate': 0.0,
                    'most_recent_file': None
                }
            
            # Sort by modification time to get most recent
            delivery_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            most_recent_file = delivery_files[0].name if delivery_files else None
            
            # Load and analyze deliveries
            total_sent = 0
            total_drafts = 0
            successful_sends = 0
            last_sent = None
            
            for file_path in delivery_files[:100]:  # Limit to last 100 for performance
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    status = data.get('response', {}).get('status')
                    if status == EmailStatus.SENT:
                        total_sent += 1
                        successful_sends += 1
                        sent_at = data.get('response', {}).get('sent_at')
                        if sent_at and not last_sent:
                            last_sent = sent_at
                    elif status == EmailStatus.DRAFT:
                        total_drafts += 1
                    elif status == EmailStatus.FAILED:
                        total_drafts += 1
                    elif status == EmailStatus.PENDING:
                        total_drafts += 1
                        
                except Exception as e:
                    logger.error(f"Error loading delivery file {file_path}: {str(e)}")
            
            total_emails = total_sent + total_drafts
            success_rate = (successful_sends / total_emails * 100) if total_emails > 0 else 0.0
            
            return {
                'total_sent': total_sent,
                'total_drafts': total_drafts,
                'last_sent': last_sent,
                'success_rate': success_rate,
                'most_recent_file': most_recent_file
            }
            
        except Exception as e:
            logger.error(f"Error getting delivery stats: {str(e)}")
            return {
                'error': str(e),
                'total_sent': 0,
                'total_drafts': 0,
                'last_sent': None,
                'success_rate': 0.0,
                'most_recent_file': None
            }
    
    def list_delivery_files(self) -> List[str]:
        """List all delivery files"""
        try:
            deliveries_dir = Path("data/deliveries")
            
            if not deliveries_dir.exists():
                return []
            
            delivery_files = [f.name for f in deliveries_dir.glob("delivery_*.json")]
            
            # Return sorted list (newest first)
            return sorted(delivery_files, reverse=True)
            
        except Exception as e:
            logger.error(f"Error listing delivery files: {str(e)}")
            return []
